"""
translation_mixin.py — 번역 실행, 용어집 추출, DOCX/HTML 출력 생성 메서드를 담은 Mixin 클래스.

TRPGTranslatorApp이 이 Mixin을 상속해 번역 탭의 모든 기능을 사용한다.
백그라운드 스레드에서 실행되는 작업은 msg_queue를 통해 GUI와 통신한다.
"""
import os
import threading
import shutil
import math
import re
import time
import itertools

import config
from file_handler import get_file_content, scan_pdf_images, detect_scanned_pdf
from translator import translate_content, refine_content, configure_genai
from korean_utils import apply_replacement

# Pre-compiled regex patterns
_RE_COLOR_TAG_SPLIT = re.compile(
    r'(<c=#[0-9a-fA-F]{6}>.*?</c>|<b>.*?</b>)',
    re.DOTALL
)
_RE_COLOR_TAG_STRIP = re.compile(
    r'<c=#[0-9a-fA-F]{6}>|</c>|<b>|</b>'
)
_RE_INVALID_FILENAME = re.compile(r'[\\/*?:"<>|]')


class TranslationMixin:
    """번역, 용어집 추출, HTML/DOCX 출력 생성 기능을 제공하는 Mixin."""

    # ── 용어집 추출 ──────────────────────────────────────────────────────

    def start_glossary_extraction(self):
        """텍스트/PDF 폴더의 내용을 읽어 Gemini API로 고유명사를 추출하고 auto_glossary.txt에 저장한다."""
        import tkinter as tk
        from tkinter import messagebox
        text_dir = self.text_dir.get()
        output_dir = self.output_dir.get()

        if not text_dir:
            messagebox.showwarning("Warning", "텍스트 폴더를 선택해주세요. (PDF 및 텍스트 파일 기반으로 추출합니다.)")
            return
        if not output_dir:
            messagebox.showwarning("Warning", "결과물 저장 폴더를 선택해주세요.")
            return

        text_files = []
        if os.path.exists(text_dir):
            for root, _, files in os.walk(text_dir):
                for f in files:
                    if f.lower().endswith(('.txt', '.pdf')):
                        text_files.append(os.path.join(root, f))

        if not text_files:
            messagebox.showwarning("Warning", "선택한 텍스트 폴더에 텍스트나 PDF 파일이 없습니다.")
            return

        self.start_btn.config(state='disabled')
        self.glossary_btn.config(state='disabled')
        self.log_area.config(state='normal')
        self.log_area.delete(1.0, tk.END)
        self.log_area.config(state='disabled')
        self.progress['value'] = 0
        self.status_label.config(text="고유명사 추출용 데이터 수집 중...")

        # Configure GenAI dynamically
        api_key = self.api_key_var.get().strip()
        model_name = self.model_name_var.get().strip()

        success, msg = configure_genai(api_key, model_name)
        if not success:
            messagebox.showerror("Config Error", f"Failed to configure AI: {msg}")
            self.start_btn.config(state='normal')
            self.glossary_btn.config(state='normal')
            return

        thread = threading.Thread(target=self.run_glossary_extraction, args=(text_files, output_dir))
        thread.daemon = True
        thread.start()

    def run_glossary_extraction(self, text_files, output_dir):
        """백그라운드 스레드에서 실행: 파일을 읽고 API로 고유명사를 추출해 저장한다."""
        from translator import extract_glossary
        import pypdf

        try:
            self.log(f"Collecting text from {len(text_files)} files for Glossary Extraction...")

            # API 토큰 소비 제한: ~30,000자 ≈ 7,500토큰으로 주요 고유명사 추출에 충분
            MAX_GLOSSARY_INPUT = 30000
            combined_text = ""
            limit_reached = False
            for i, filepath in enumerate(text_files):
                if limit_reached:
                    break
                self.update_progress(i, len(text_files), f"Reading: {os.path.basename(filepath)}")

                try:
                    if filepath.lower().endswith('.pdf'):
                        with open(filepath, 'rb') as f:
                            reader = pypdf.PdfReader(f)
                            for page in reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    combined_text += page_text + "\n"
                                if len(combined_text) >= MAX_GLOSSARY_INPUT:
                                    limit_reached = True
                                    break
                    else:
                        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                            combined_text += f.read() + "\n"
                        if len(combined_text) >= MAX_GLOSSARY_INPUT:
                            limit_reached = True
                except Exception as e:
                    self.log(f"Error reading {filepath}: {e}")

            if limit_reached:
                combined_text = combined_text[:MAX_GLOSSARY_INPUT]
                self.log(f"(입력 크기 제한 {MAX_GLOSSARY_INPUT:,}자 도달. 일부 파일은 제외됨)")

            if not combined_text.strip():
                self.msg_queue.put(("error", "추출할 텍스트 내용이 없습니다."))
                self.msg_queue.put(("done_glossary", None))
                return

            self.log("Sending collected text to AI for Glossary Extraction...")
            self.update_progress(len(text_files), len(text_files), "AI 고유명사 추출 중... (잠시만 기다려주세요)")

            glossary_result = ""
            for chunk in extract_glossary(combined_text):
                if chunk.startswith("Error:"):
                    self.log(chunk)
                    self.msg_queue.put(("error", f"사전 추출 실패: {chunk}"))
                    self.msg_queue.put(("done_glossary", None))
                    return

                glossary_result += chunk

            if not glossary_result.strip():
                self.log("AI returned empty glossary or extraction failed.")
                self.msg_queue.put(("error", "사전 추출에 실패했거나 추출된 단어가 없습니다."))
                self.msg_queue.put(("done_glossary", None))
                return

            # Write out to output_dir
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            out_file = os.path.join(output_dir, "auto_glossary.txt")
            with open(out_file, 'w', encoding='utf-8') as f:
                f.write(glossary_result)

            self.log(f"Glossary successfully saved to: {out_file}")

            # Queue success message so the main thread shows a nice popup
            # And we set the trans_glossary_file variable to this new file automatically!
            self.msg_queue.put(("glossary_success", out_file))

        except Exception as e:
            self.log(f"Fatal error during formatting: {str(e)}")
            self.msg_queue.put(("error", f"Error: {str(e)}"))
        finally:
            self.msg_queue.put(("done_glossary", None))

    # ── 번역 시작 / 실행 ─────────────────────────────────────────────────

    def start_translation(self):
        """번역 실행을 준비한다: 파일 수집, PDF 이미지 선택 다이얼로그, Gemini 설정 후 백그라운드 스레드를 시작한다."""
        import tkinter as tk
        from tkinter import messagebox
        from dialogs import ImageSelectionDialog

        text_dir = self.text_dir.get()
        image_dir = self.image_dir.get()
        glossary_file = self.trans_glossary_file.get()

        if not text_dir and not image_dir:
            messagebox.showwarning("Warning", "텍스트 폴더 또는 이미지 폴더를 선택해주세요.")
            return

        text_files = []
        if text_dir and os.path.exists(text_dir):
            for root, _, files in os.walk(text_dir):
                for f in files:
                    if f.lower().endswith(('.txt', '.pdf')):
                        rel_dir = os.path.relpath(root, text_dir)
                        if rel_dir == '.':
                            rel_dir = ''
                        text_files.append((os.path.join(root, f), rel_dir))

        image_files = []
        if image_dir and os.path.exists(image_dir):
            for root, _, files in os.walk(image_dir):
                for f in files:
                    if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp')):
                        rel_dir = os.path.relpath(root, image_dir)
                        if rel_dir == '.':
                            rel_dir = ''
                        image_files.append((os.path.join(root, f), rel_dir))

        if not text_files and not image_files:
            messagebox.showwarning("Warning", "선택한 폴더에 처리할 파일이 없습니다.")
            return

        self.start_btn.config(state='disabled')
        self.log_area.config(state='normal')
        self.log_area.delete(1.0, tk.END)
        self.log_area.config(state='disabled')
        self.progress['value'] = 0
        self.status_label.config(text="준비 중...")

        # --- Pre-scan PDF Images for User Selection ---
        allowed_image_ids = None  # Default: Allow all (if None)

        pdf_files = [f[0] for f in text_files if f[0].lower().endswith('.pdf')]
        if pdf_files:
            self.status_label.config(text="PDF 이미지 스캔 중...")
            try:
                self.root.update()
            except Exception:
                pass

            img_selection_map = {}  # filepath -> set(allowed_ids)
            has_images = False

            for pdf_path in pdf_files:
                fname = os.path.basename(pdf_path)
                self.log(f"Scanning PDF for images: {fname}")

                # ── PDF 스캔본 감지 ──────────────────────────────────────
                scanned_pages = detect_scanned_pdf(pdf_path)
                if scanned_pages:
                    messagebox.showwarning(
                        "스캔 PDF 감지",
                        f"{fname}\n텍스트 레이어 없는 페이지: {scanned_pages}\n\n"
                        "해당 페이지는 번역 결과가 없거나 이미지만 추출될 수 있습니다."
                    )

                try:
                    found = scan_pdf_images(pdf_path)
                    if found:
                        self.log(f"  - Found {len(found)} images in {fname}. Opening dialog...")
                        # Show dialog for THIS file (must run in main thread)
                        dlg = ImageSelectionDialog(self.root, found, fname)
                        self.root.wait_window(dlg)
                        img_selection_map[pdf_path] = dlg.selected_ids
                        has_images = True
                    else:
                        self.log(f"  - No images extracted from {fname}.")
                except Exception as e:
                    self.log(f"  - Error scanning {fname}: {str(e)}")
                    messagebox.showerror("Scan Error", f"Error scanning PDF {fname}:\n{str(e)}")

            if has_images:
                allowed_image_ids = img_selection_map
            else:
                self.log("No images found in any PDF files.")

        # Configure GenAI dynamically
        api_key = self.api_key_var.get().strip()
        model_name = self.model_name_var.get().strip()

        self.log(f"Using Model: {model_name}")

        # Parse Glossary File
        glossary = self.load_glossary_from_file(glossary_file)

        success, msg = configure_genai(api_key, model_name, glossary=glossary)
        self.log(f"AI 초기화: {msg}")
        if not success:
            messagebox.showerror("Config Error", f"Failed to configure AI: {msg}\nPlease check Settings tab.")
            self.start_btn.config(state='normal')
            if hasattr(self, 'glossary_btn') and self.glossary_btn:
                self.glossary_btn.config(state='normal')
            return

        output_dir = self.output_dir.get()
        if not output_dir:
            messagebox.showwarning("Warning", "출력 폴더를 선택해주세요.")
            self.start_btn.config(state='normal')
            return

        # Save these paths to settings immediately
        config.save_settings(
            self.settings.get("api_key", ""),
            self.settings.get("model_name", ""),
            translation_rules=self.settings.get("translation_rules", config.DEFAULT_TRANSLATION_RULES),
            system_prompt=self.settings.get("system_prompt", config.DEFAULT_SYSTEM_PROMPT),
            last_text_dir=text_dir,
            last_image_dir=image_dir,
            last_output_dir=output_dir,
            last_trans_glossary=glossary_file,
            docx_font_name=self.docx_font_name.get()
        )

        # Update self.settings in memory as well
        self.settings["last_text_dir"] = text_dir
        self.settings["last_image_dir"] = image_dir
        self.settings["last_output_dir"] = output_dir
        self.settings["last_trans_glossary"] = glossary_file

        # Start Thread
        thread = threading.Thread(target=self.run_translation, args=(
            text_files, image_files, glossary, output_dir,
            self.refine_enabled.get(), allowed_image_ids,
            self.docx_output_enabled.get(), self.docx_font_name.get(),
            self.docx_remove_headers.get(),
            self.resume_enabled.get()
        ))
        thread.daemon = True
        thread.start()

    def run_translation(self, text_files, image_files, glossary, output_dir, refine_enabled,
                        allowed_image_ids_map=None, docx_enabled=False, docx_font_name="바탕",
                        docx_remove_headers=False, resume_enabled=False):
        """백그라운드 스레드에서 실행: 파일 단위로 번역을 수행하고 HTML(및 선택적으로 DOCX)로 저장한다.

        allowed_image_ids_map: {filepath: set(image_ids)} — None이면 모든 이미지 번역.
        resume_enabled: True이면 이미 출력 파일이 존재하는 경우 해당 파일을 건너뛴다.
        """
        start_time = time.time()
        stats = {"files": 0, "chunks": 0, "errors": 0}

        try:
            files_to_process = []
            files_to_process.extend(text_files)
            files_to_process.extend(image_files)

            files_to_process.sort(key=lambda x: x[0])
            total_files = len(files_to_process)

            if total_files == 0:
                self.msg_queue.put(("error", "No files found to process."))
                return

            self.log(f"Found {total_files} files.")

            # Create output directory if not exists
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # Create images subdirectory for extracted PDF images
            extracted_images_dir = os.path.join(output_dir, "images")
            if not os.path.exists(extracted_images_dir):
                os.makedirs(extracted_images_dir)

            # Load Translation Rules
            rules = config.DEFAULT_TRANSLATION_RULES  # noqa: F841

            # 1. Process Text/PDF Files
            for i, (filepath, rel_dir) in enumerate(files_to_process):
                filename = str(os.path.basename(filepath))
                self.update_progress(i, total_files, f"Processing: {filename}...")
                self.log(f"Starting: {filename}")

                base_name = os.path.splitext(filename)[0]
                output_filename = f"{base_name}_translated.html"

                current_output_dir = os.path.join(output_dir, rel_dir) if rel_dir else output_dir
                if not os.path.exists(current_output_dir):
                    os.makedirs(current_output_dir)

                # Check for Image (Direct Image File)
                ext = str(os.path.splitext(filepath)[1]).lower()
                if ext in ['.png', '.jpg', '.jpeg', '.webp']:
                    # 이미지 파일은 Section 2(이미지 합본)에서 전담 처리 — 여기서는 건너뜀
                    continue
                else:
                    # Text/PDF files go to main output dir
                    output_path = os.path.join(current_output_dir, output_filename)

                # Resume: skip already-translated files
                if resume_enabled and os.path.exists(output_path):
                    self.log(f"이미 번역됨, 건너뜀: {filename}")
                    continue

                # Refined output preparation
                refined_outfile = None
                if refine_enabled:
                    if ext in ['.png', '.jpg', '.jpeg', '.webp']:
                        refined_dir = os.path.join(current_output_dir, "individual_images", "refined")
                    else:
                        refined_dir = os.path.join(current_output_dir, "refined")

                    if not os.path.exists(refined_dir):
                        os.makedirs(refined_dir)
                    refined_path = os.path.join(refined_dir, f"{base_name}_polished.html")
                    refined_outfile = open(refined_path, 'w', encoding='utf-8')
                    # Write Header to Refined
                    refined_outfile.write("""
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Polished: """ + filename + """</title>
    <style>
        body { font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', sans-serif; background-color: #f0f2f5; padding: 20px; margin: 0; }
        .container { max-width: 1000px; margin: 0 auto; }
        .card { background: white; padding: 25px; margin-bottom: 20px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }
        .card-header { font-weight: bold; color: #2c3e50; font-size: 1.1em; margin-bottom: 15px; border-bottom: 2px solid #eee; padding-bottom: 8px; }
        .card-content { display: flex; flex-wrap: wrap; gap: 20px; align-items: flex-start; }
        .text-content { white-space: pre-wrap; line-height: 1.8; color: #333; font-size: 1.05em; flex: 2 1 400px; min-width: 0; }
        .image-content { flex: 1 1 300px; max-width: 100%; }
        .image-content img { max-width: 100%; height: auto; display: block; border-radius: 8px; border: 1px solid #ddd; }
        @media (max-width: 600px) {
            .card-content { flex-direction: column; }
            .image-content, .text-content { flex: 1 1 auto; max-width: 100%; }
                        }
    </style>
</head>
<body>
    <div class="container">
        <h1>""" + filename + """ (교열본)</h1>
""")

                # Initialize DOCX Document if enabled
                docx_doc = None
                if docx_enabled:
                    try:
                        import docx
                        from docx.oxml.ns import qn
                        docx_doc = docx.Document()

                        # Set default font to user-defined font for East Asian
                        style = docx_doc.styles['Normal']
                        style.font.name = docx_font_name
                        style.element.rPr.rFonts.set(qn('w:eastAsia'), docx_font_name)

                    except ImportError:
                        self.log("python-docx not installed. Skipping DOCX creation.")
                        docx_enabled = False

                try:
                    with open(output_path, 'w', encoding='utf-8') as outfile:
                        outfile.write("""
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Translation: """ + filename + """</title>
    <style>
        body { font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', sans-serif; background-color: #f0f2f5; padding: 20px; margin: 0; }
        .container { max-width: 1000px; margin: 0 auto; }
        .card { background: white; padding: 25px; margin-bottom: 20px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }
        .card-header { font-weight: bold; color: #2c3e50; font-size: 1.1em; margin-bottom: 15px; border-bottom: 2px solid #eee; padding-bottom: 8px; }
        .card-content { display: flex; flex-wrap: wrap; gap: 20px; align-items: flex-start; }
        .text-content { white-space: pre-wrap; line-height: 1.8; color: #333; font-size: 1.05em; flex: 2 1 400px; min-width: 0; }
        .image-content { flex: 1 1 300px; max-width: 100%; }
        .image-content img { max-width: 100%; height: auto; display: block; border-radius: 8px; border: 1px solid #ddd; }
        @media (max-width: 600px) {
            .card-content { flex-direction: column; }
            .image-content, .text-content { flex: 1 1 auto; max-width: 100%; }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>""" + filename + """ 번역 결과</h1>
""")
                        # Get specific allowed IDs for this file, if any
                        file_allowed_ids = None
                        if allowed_image_ids_map and filepath in allowed_image_ids_map:
                            file_allowed_ids = allowed_image_ids_map[filepath]

                        known_colors = {}

                        # Pre-create images directory once before the chunk loop
                        current_extracted_images_dir = os.path.join(current_output_dir, "images")
                        os.makedirs(current_extracted_images_dir, exist_ok=True)

                        for header, content, content_type in get_file_content(
                                filepath, allowed_image_ids=file_allowed_ids, extract_color=docx_enabled):
                            self.log(f"  - Translating: {header}")
                            if content_type == 'text' and isinstance(content, str):
                                self.log(f"    ⏳ {len(content)}자 번역 중...")

                            translated_text = translate_content(content, content_type, log_fn=self.log)

                            # 번역 오류 시 출력 파일 삭제 후 중단 (건너뛰기가 오인하지 않도록)
                            if isinstance(translated_text, str) and (
                                translated_text.startswith("번역 중 오류 발생:") or
                                translated_text.startswith("번역 실패:")
                            ):
                                self.log(f"  ❌ 번역 오류로 중단: {translated_text}")
                                raise Exception(translated_text)

                            # Refine if enabled
                            refined_text_result = None
                            if refine_enabled:
                                self.log(f"    > Polishing...")
                                refined_text_result = refine_content(translated_text)

                            # Add to DOCX if requested
                            if docx_doc is not None:
                                if content_type != 'pdf_image':
                                    p = docx_doc.add_paragraph()
                                    # Header (optional)
                                    if not docx_remove_headers:
                                        p.add_run(f"[{header}]\n").bold = True

                                    # Split by internal tags: <c=#RRGGBB>...</c> and <b>...</b>
                                    parts = _RE_COLOR_TAG_SPLIT.split(translated_text)
                                    for part in parts:
                                        if part.startswith('<b>') and part.endswith('</b>'):
                                            # Bold tag — may contain nested <c=> tag
                                            inner = part[3:-4]
                                            if inner.startswith('<c=') and inner.endswith('</c>'):
                                                color_hex = inner[3:10].lower()
                                                inner_text = inner[11:-4]
                                            else:
                                                color_hex = None
                                                inner_text = inner
                                            run = p.add_run(inner_text)
                                            run.bold = True
                                            if color_hex:
                                                try:
                                                    from docx.shared import RGBColor
                                                    r, g, b = int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16)
                                                    if r > 240 and g > 240 and b > 240:
                                                        r, g, b = 0, 0, 0
                                                    run.font.color.rgb = RGBColor(r, g, b)
                                                except Exception:
                                                    pass
                                        elif part.startswith('<c=') and part.endswith('</c>'):
                                            color_hex = part[3:10].lower()
                                            inner_text = part[11:-4]
                                            run = p.add_run(inner_text)
                                            try:
                                                from docx.shared import RGBColor

                                                r, g, b = int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16)
                                                assigned_hex = color_hex

                                                # 색상 스냅: 유사한 색(유클리드 거리 < 60)은 동일 색으로 통일해
                                                # PDF 스팬마다 미세하게 다른 RGB가 여러 색으로 분리되는 현상을 방지
                                                for k_hex, (kr, kg, kb) in known_colors.items():
                                                    dist = math.sqrt((r - kr) ** 2 + (g - kg) ** 2 + (b - kb) ** 2)
                                                    if dist < 60:
                                                        assigned_hex = k_hex
                                                        r, g, b = kr, kg, kb
                                                        break

                                                if assigned_hex not in known_colors:
                                                    known_colors[assigned_hex] = (r, g, b)

                                                # 거의 흰색(RGB 각 채널 > 240)은 DOCX 흰 배경에서 보이지 않으므로 검정으로 변환
                                                if r > 240 and g > 240 and b > 240:
                                                    r, g, b = 0, 0, 0

                                                run.font.color.rgb = RGBColor(r, g, b)
                                            except Exception:
                                                pass
                                        elif part:
                                            p.add_run(part)

                            # Strip tags for HTML output
                            safe_text_for_html = _RE_COLOR_TAG_STRIP.sub('', translated_text)
                            safe_header = header.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            safe_text = safe_text_for_html.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                            html_chunk = ""

                            if content_type == 'pdf_image':
                                # Save the image
                                img_filename = f"{base_name}_{header.replace('[', '').replace(']', '').replace(' ', '_')}.png"
                                # Clean filename
                                img_filename = _RE_INVALID_FILENAME.sub("", img_filename)

                                img_save_path = os.path.join(current_extracted_images_dir, img_filename)
                                content.save(img_save_path)

                                # Relative path for HTML
                                rel_img_path = f"images/{img_filename}"

                                html_chunk = f"""
        <div class="card">
            <div class="card-header">{safe_header}</div>
            <div class="card-content">
                <div class="image-content"><img src="{rel_img_path}" alt="PDF Image"></div>
                <div class="text-content">{safe_text}</div>
            </div>
        </div>
"""
                            else:
                                html_chunk = f"""
        <div class="card">
            <div class="card-header">{safe_header}</div>
            <div class="card-content">
                <div class="text-content">{safe_text}</div>
            </div>
        </div>
"""
                            outfile.write(html_chunk)

                            # Write to Refined Output
                            if refined_outfile and refined_text_result:
                                safe_refined = refined_text_result.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                                refined_chunk = html_chunk.replace(safe_text, safe_refined)
                                refined_outfile.write(refined_chunk)

                            preview = f"[{header}]\n{translated_text[:100]}...\n\n"
                            self.log(preview)
                            stats["chunks"] += 1

                        outfile.write("\n    </div>\n</body>\n</html>")
                        if refined_outfile:
                            refined_outfile.write("\n    </div>\n</body>\n</html>")
                            refined_outfile.close()

                        if docx_doc is not None:
                            docx_path = os.path.join(current_output_dir, f"{base_name}_translated.docx")
                            docx_doc.save(docx_path)
                            self.log(f"Saved DOCX: {os.path.basename(docx_path)}")

                    stats["files"] += 1
                except Exception as e:
                    self.log(f"Error processing {filename}: {str(e)}")
                    # Print full traceback for debugging safely
                    import traceback
                    traceback.print_exc()
                    stats["errors"] += 1
                    # 불완전한 출력 파일 삭제 — 건너뛰기 기능이 오인하지 않도록
                    try:
                        if 'output_path' in locals() and os.path.exists(output_path):
                            os.remove(output_path)
                            self.log(f"  (불완전한 출력 파일 삭제됨: {os.path.basename(output_path)})")
                    except Exception:
                        pass

            # 2. Process Image Files (Combined HTML)
            if image_files:
                # Group image files by relative directory
                image_files_by_dir: dict[str, list[str]] = {}
                for filepath, rel_dir in image_files:
                    if rel_dir not in image_files_by_dir:
                        image_files_by_dir[rel_dir] = []
                    image_files_by_dir[rel_dir].append(filepath)

                for rel_dir, imgs in image_files_by_dir.items():
                    current_output_dir = os.path.join(output_dir, rel_dir) if rel_dir else output_dir
                    if not os.path.exists(current_output_dir):
                        os.makedirs(current_output_dir)

                    output_filename = "images_translated.html"
                    output_path = os.path.join(current_output_dir, output_filename)

                    current_extracted_images_dir = os.path.join(current_output_dir, "images")
                    if not os.path.exists(current_extracted_images_dir):
                        os.makedirs(current_extracted_images_dir)

                    msg_dir = rel_dir if rel_dir else "root"
                    self.log(f"Starting Images: Consolidating {len(imgs)} images into {output_filename} in {msg_dir}")

                    try:
                        with open(output_path, 'w', encoding='utf-8') as outfile:
                            outfile.write("""
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Image Translations</title>
    <style>
        body { font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', sans-serif; background-color: #f0f2f5; padding: 20px; margin: 0; }
        .container { max-width: 1000px; margin: 0 auto; }
        .card { background: white; padding: 25px; margin-bottom: 20px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }
        .card-header { font-weight: bold; color: #2c3e50; font-size: 1.1em; margin-bottom: 15px; border-bottom: 2px solid #eee; padding-bottom: 8px; }
        .card-content { display: flex; flex-wrap: wrap; gap: 20px; align-items: flex-start; }
        .text-content { white-space: pre-wrap; line-height: 1.8; color: #333; font-size: 1.05em; flex: 2 1 400px; min-width: 0; }
        .image-content { flex: 1 1 300px; max-width: 100%; }
        .image-content img { max-width: 100%; height: auto; display: block; border-radius: 8px; border: 1px solid #ddd; }
        @media (max-width: 600px) {
            .card-content { flex-direction: column; }
            .image-content, .text-content { flex: 1 1 auto; max-width: 100%; }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>이미지 파일 번역 결과 모음</h1>
""")
                            for i, filepath in enumerate(imgs):
                                filename = os.path.basename(filepath)

                                current_progress_idx = i
                                self.update_progress(current_progress_idx, total_files, f"Processing Image: {filename}...")

                                # Standard Images: Copy to output folder for display
                                try:
                                    img_save_path = os.path.join(current_extracted_images_dir, filename)
                                    shutil.copy2(filepath, img_save_path)
                                    rel_img_path = f"images/{filename}"
                                except Exception as copy_err:
                                    self.log(f"Error copying image {filename}: {copy_err}")
                                    rel_img_path = ""

                                for header, content, content_type in get_file_content(filepath):
                                    self.log(f"  - Translating image: {header}")
                                    translated_text = translate_content(content, content_type, log_fn=self.log)

                                    # Apply Glossary
                                    for original, translated in glossary.items():
                                        translated_text = apply_replacement(translated_text, original, translated)

                                    safe_header = header.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                                    safe_content = translated_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                                    if rel_img_path:
                                        html_chunk = f"""
        <div class="card">
            <div class="card-header">{safe_header}</div>
            <div class="card-content">
                 <div class="image-content"><img src="{rel_img_path}" alt="Image Preview"></div>
                 <div class="text-content">{safe_content}</div>
            </div>
        </div>
"""
                                    else:
                                        html_chunk = f"""
        <div class="card">
            <div class="card-header">{safe_header}</div>
            <div class="card-content">{safe_content}</div>
        </div>
"""
                                    outfile.write(html_chunk)
                                    preview = f"[{header}]\n{translated_text[:100]}...\n\n"
                                    self.log(preview)

                            outfile.write("\n    </div>\n</body>\n</html>")

                    except Exception as e:
                        self.log(f"Error processing images in {msg_dir}: {str(e)}")

            self.update_progress(total_files, total_files, "Done")
            stats["elapsed"] = time.time() - start_time
            self.msg_queue.put(("done", stats))

        except Exception as e:
            self.msg_queue.put(("error", f"Critical Error: {str(e)}"))

    # ── 1페이지 미리번역 (Quick Preview) ─────────────────────────────────────

    def start_quick_preview(self):
        """텍스트/PDF 파일 첫 3청크를 번역해 팝업 창에 표시한다."""
        import tkinter as tk
        from tkinter import messagebox

        text_dir = self.text_dir.get()
        if not text_dir:
            messagebox.showwarning("Warning", "텍스트 폴더를 선택해주세요.")
            return

        # Collect first text/PDF file
        preview_file = None
        if os.path.exists(text_dir):
            for root, _, files in os.walk(text_dir):
                for f in sorted(files):
                    if f.lower().endswith(('.txt', '.pdf')):
                        preview_file = os.path.join(root, f)
                        break
                if preview_file:
                    break

        if not preview_file:
            messagebox.showwarning("Warning", "텍스트 또는 PDF 파일을 찾을 수 없습니다.")
            return

        api_key = self.api_key_var.get().strip()
        model_name = self.model_name_var.get().strip()
        glossary_file = self.trans_glossary_file.get().strip() if hasattr(self, 'trans_glossary_file') else ""
        glossary = self.load_glossary_from_file(glossary_file)

        success, msg = configure_genai(api_key, model_name, glossary=glossary)
        if not success:
            messagebox.showerror("Config Error", f"AI 설정 실패:\n{msg}")
            return

        filename = os.path.basename(preview_file)

        def _run():
            result_parts = []
            try:
                chunks = itertools.islice(get_file_content(preview_file), 3)
                for header, content, content_type in chunks:
                    if content_type in ('pdf_image', 'pdf_image_embed', 'image'):
                        continue  # Skip images in quick preview
                    translated = translate_content(content, content_type, log_fn=self.log)
                    result_parts.append(f"=== {header} ===\n{translated}")
            except Exception as e:
                result_parts.append(f"[오류] {str(e)}")

            result_text = "\n\n".join(result_parts) if result_parts else "(번역 결과 없음)"
            self.msg_queue.put(("preview_result", (filename, result_text)))

        thread = threading.Thread(target=_run, daemon=True)
        thread.start()
        self.log(f"미리보기 번역 시작: {filename} (최대 3청크)")
